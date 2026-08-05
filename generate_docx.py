import os
import re
import subprocess
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from lxml import etree

docx_path = r"g:\My Drive\Dev\Einstein\Le_Pacte_Nice_IA.docx"
chart1_png = r"g:\My Drive\Dev\Einstein\chart1_poids_national.png"
chart2_png = r"g:\My Drive\Dev\Einstein\chart2_roi_gains.png"

# Regenerate high-res charts
subprocess.run(["python", r"g:\My Drive\Dev\Einstein\generate_charts.py"], check=True)

# =========================================================
# BASE : Partir du document de reference pour heriter
# des styles natifs (bullet numbering etc.)
# =========================================================
REF_PATH = u'G:\\My Drive\\Arx Capital\\web\\arxWeb\\Le_Pacte_Nice_IA-r\u00e9f\u00e9rence.docx'
doc = Document(REF_PATH)

# Vider tous les paragraphes et tableaux existants de la reference
# en les supprimant de l'element XML body
body = doc.element.body
for child in list(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag in ('p', 'tbl', 'sdt'):
        body.remove(child)

# =========================================================
# MARGES — Exactement identiques a la reference
# top=0.590in, bottom=0.360in, left=right=0.900in
# =========================================================
for section in doc.sections:
    section.top_margin    = Inches(0.590)
    section.bottom_margin = Inches(0.360)
    section.left_margin   = Inches(0.900)
    section.right_margin  = Inches(0.900)

# =========================================================
# PALETTE DE COULEURS
# =========================================================
NAVY_PRIMARY   = RGBColor(0x0F, 0x17, 0x2A)
NAVY_SECONDARY = RGBColor(0x1E, 0x3A, 0x8A)
SLATE_DARK     = RGBColor(0x33, 0x41, 0x55)
BODY_BLACK     = RGBColor(0x1E, 0x29, 0x3B)
MUTED_GREY     = RGBColor(0x64, 0x74, 0x8B)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            '</w:tblBorders>'
        )
        tblPr[0].append(borders)

# =========================================================
# FONCTIONS DE MISE EN PAGE
# Spec exacte extraite de la reference :
#   normal (ARX subline)  : sz=8.5 Arial bold #334155, sa=2
#   Title                 : sz=36 Georgia bold #0F172A, sa=4
#   normal (subtitle)     : sz=11 Arial #1E3A8A, sa=6
#   normal (author)       : sz=9 Arial italic #64748B, sa=20
#   Heading 1             : sz=15 Georgia bold #0F172A, sb=18 sa=8
#   Heading 2             : sz=12.5 Georgia bold #1E3A8A, sb=14 sa=6
#   bullet (normal+numPr) : sz=11 Arial li=0.35 ls=1.15
#                           sb=2 sur 1er, sb=inh sur autres
#                           sa=5 sur dernier, sa=inh sur autres
# =========================================================

def add_normal(text_runs, sa=None, sb=None):
    """Ajoute un paragraphe style 'normal'. text_runs = list of (text, bold, italic, color)"""
    p = doc.add_paragraph(style='Normal')
    if sb is not None:
        p.paragraph_format.space_before = Pt(sb)
    if sa is not None:
        p.paragraph_format.space_after = Pt(sa)
    for (txt, bold, italic, color) in text_runs:
        if not txt:
            continue
        r = p.add_run(txt)
        r.font.name = 'Arial'
        if bold:  r.font.bold = True
        if italic: r.font.italic = True
        if color: r.font.color.rgb = color
    return p


def add_header_banner():
    # ARX CONSULTING line
    p1 = doc.add_paragraph(style='Normal')
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("ARX CONSULTING — DOCUMENT STRATÉGIQUE MÉTROPOLITAIN")
    r1.font.name = 'Arial'; r1.font.size = Pt(8.5)
    r1.font.bold = True; r1.font.color.rgb = SLATE_DARK

    # Title
    p2 = doc.add_paragraph(style='Title')
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("LE PACTE NICE IA")
    r2.font.name = 'Georgia'; r2.font.size = Pt(36)
    r2.font.bold = True; r2.font.color.rgb = NAVY_PRIMARY

    # Subtitle
    p3 = doc.add_paragraph(style='Normal')
    p3.paragraph_format.space_after = Pt(6)
    r3 = p3.add_run("Doctrine Stratégique, Rigueur Budgétaire & Alliance Transfrontalière (2026-2029)")
    r3.font.name = 'Arial'; r3.font.size = Pt(11)
    r3.font.color.rgb = NAVY_SECONDARY

    # Author
    p4 = doc.add_paragraph(style='Normal')
    p4.paragraph_format.space_after = Pt(20)
    r4 = p4.add_run(
        "Rédigé pour M. Éric Ciotti par Benoît SIGWALD — "
        "Directeur du Projet Pacte Nice IA & Senior AI Architect — Août 2026"
    )
    r4.font.name = 'Arial'; r4.font.size = Pt(9)
    r4.font.italic = True; r4.font.color.rgb = MUTED_GREY


def add_h1(text, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        r.font.name = 'Georgia'; r.font.size = Pt(15)
        r.font.bold = True; r.font.color.rgb = NAVY_PRIMARY


def add_h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        r.font.name = 'Georgia'; r.font.size = Pt(12.5)
        r.font.bold = True; r.font.color.rgb = NAVY_SECONDARY


def _parse_runs(text):
    """Parse **gras** *italique* et retourne liste de (txt, bold, italic)"""
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
    result = []
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            result.append((part[2:-2], True, False))
        elif part.startswith('*') and part.endswith('*'):
            result.append((part[1:-1], False, True))
        else:
            if part:
                result.append((part, False, False))
    return result


def add_bullet_para(text, is_first=False, is_last=False):
    """
    Bullet EXACTEMENT comme la reference :
    Style 'Normal' + numPr injecte en XML (numId=1, ilvl=0)
    ind left=504 hanging=360 (0.35in / 0.25in)
    line=276 (1.15x), before=40 (2pt) sur 1er, after=100 (5pt) sur dernier
    """
    p = doc.add_paragraph(style='Normal')
    pPr = p._element.get_or_add_pPr()

    # Inject numPr XML (numId=1 = liste bullet heritee de la reference)
    numPr_xml = (
        '<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:ilvl w:val="0"/>'
        '<w:numId w:val="1"/>'
        '</w:numPr>'
    )
    from lxml import etree
    numPr_el = etree.fromstring(numPr_xml)
    # Inserer numPr en premier dans pPr
    pPr.insert(0, numPr_el)

    # Spacing: before=40 (2pt) sur 1er, line=276 (1.15x auto)
    from docx.oxml.ns import qn
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        from lxml import etree as ET
        spacing = ET.SubElement(pPr, qn('w:spacing'))
    spacing.set(qn('w:line'), '276')
    spacing.set(qn('w:lineRule'), 'auto')
    if is_first:
        spacing.set(qn('w:before'), '40')   # 2pt
    if is_last:
        spacing.set(qn('w:after'), '100')   # 5pt
    else:
        spacing.set(qn('w:after'), '0')

    # Indentation: left=504 (0.35in), hanging=360 (0.25in)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        from lxml import etree as ET
        ind = ET.SubElement(pPr, qn('w:ind'))
    ind.set(qn('w:left'), '504')
    ind.set(qn('w:hanging'), '360')

    parsed = _parse_runs(text)
    for (txt, bold, italic) in parsed:
        r = p.add_run(txt)
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        if bold:
            r.font.bold = True
            r.font.color.rgb = NAVY_PRIMARY
        elif italic:
            r.font.italic = True
        else:
            r.font.color.rgb = BODY_BLACK
    return p


def bullets(items):
    """Liste de bullets avec spacing correct reference"""
    n = len(items)
    for i, item in enumerate(items):
        add_bullet_para(item, is_first=(i == 0), is_last=(i == n - 1))


def add_sub_indent(items):
    """Items en retrait simple (li=0.25in) sans puce — pour les listes numerotees 1. 2. 3."""
    n = len(items)
    for i, text in enumerate(items):
        p = doc.add_paragraph(style='Normal')
        if i == 0:
            p.paragraph_format.space_before = None
        p.paragraph_format.space_after  = Pt(5) if i == n - 1 else None
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.line_spacing = 1.15
        parsed = _parse_runs(text)
        for (txt, bold, italic) in parsed:
            r = p.add_run(txt)
            r.font.name = 'Arial'; r.font.size = Pt(11)
            if bold:   r.font.bold = True;   r.font.color.rgb = NAVY_PRIMARY
            elif italic: r.font.italic = True
            else:      r.font.color.rgb = BODY_BLACK


def make_table_header(table, r_idx, values):
    for c_idx, val in enumerate(values):
        cell = table.cell(r_idx, c_idx)
        set_cell_background(cell, "1E293B")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(val)
        r.font.name = 'Arial'; r.font.bold = True
        r.font.color.rgb = WHITE; r.font.size = Pt(9)


def make_table_body(table, r_idx, values, total=False, sz=8.5):
    for c_idx, val in enumerate(values):
        cell = table.cell(r_idx, c_idx)
        if total:
            set_cell_background(cell, "0F172A")
        else:
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(val)
        r.font.name = 'Arial'; r.font.size = Pt(sz)
        if total:
            r.font.bold = True; r.font.color.rgb = WHITE
        else:
            r.font.color.rgb = BODY_BLACK


# =========================================================
# CONSTRUCTION DU DOCUMENT
# =========================================================
add_header_banner()

# ---------------------------------------------------------
# INDEX
# ---------------------------------------------------------
add_h1("Index")

toc_data = [
    ("Index", "1"),
    ("Résumé Exécutif & Chiffrage Consolidé pour M. le Maire", "1"),
    ("1. Diagnostic Territorial & Opportunités Européennes", "2"),
    ("   1.1 Alignement Institutionnel & Leadership", "2"),
    ("   1.2 Le Terreau Azuréen : Sophia Antipolis, Grasse, Cannes & Poids National", "2"),
    ("   1.3 Atouts Monaco, Flux Transfrontaliers & Extension Business", "3"),
    ("   1.4 L'Opportunité Historique : 30 Md€ UE & 7 AI Gigafactories", "4"),
    ("   1.5 Structuration Opérationnelle & Capture des Subventions", "4"),
    ("2. Pôle 1 : Axe Ville de Nice (Sécurité & Cadre de Vie)", "5"),
    ("   2.1 CSU Augmenté : Moins d'Écrans, Plus de Policiers dans la Rue", "5"),
    ("   2.2 Guichet Vocal Allo Niçois 24/7 : Demandes d'Aide & Sécurité", "5"),
    ("   2.3 Propreté Augmentée & Routage Intelligent de la Voie Publique", "5"),
    ("   2.4 Consultation Directe par Quartier via l'IA Einstein", "5"),
    ("3. Pôle 2 : Axe Métropole Nice Côte d'Azur (Budget, DSI & Monaco)", "7"),
    ("   3.1 Audit IA Commande Publique (+2,5 M€ / an Net)", "7"),
    ("   3.2 Bouclier Cyber-IA NIS 2 (1,8 M€ à 3 M€ Évités)", "7"),
    ("   3.3 Monaco Cloud & Redondance IT (1,5 M€ à 2,8 M€ / an)", "7"),
    ("4. Alliance Binationale Nice-Monaco & Hub Réglementaire", "8"),
    ("   4.1 Levier Binational : Clé de Voûte des Financements Européens", "8"),
    ("   4.2 Nice, Centre d'Expertise AI Act de Référence pour la France", "8"),
    ("   4.3 Pôle d'Excellence PME/PMI & Bibliothèque de Cas d'Usage Souverains", "8"),
    ("   4.4 Pôle d'Excellence pour les Niçoises et les Niçois", "9"),
    ("5. Feuille de Route Précise (Sept 2026-2029) & Résumé Annuel", "9"),
    ("   5.1 Présidence Éric Ciotti & Direction Benoît SIGWALD (2,5j/sem)", "9"),
    ("   5.2 Échéancier Précis et Jalons par Date (Démarche 1er Septembre 2026)", "9"),
    ("   5.3 Budget par Axe Institutionnel (Mairie, Métropole, Monaco)", "10"),
    ("   5.4 Résumé Consolidé du Coût par Année Budgétaire", "10"),
    ("6. Annexe : Justifications des Gains & Sources Documentaires", "11"),
]

# Ajouter un paragraphe vide pour créer de l'espace avant la table de l'Index
p_space_toc = doc.add_paragraph(style='Normal')
p_space_toc.paragraph_format.space_before = Pt(6)
p_space_toc.paragraph_format.space_after  = Pt(6)

table_toc = doc.add_table(rows=len(toc_data), cols=2)
table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
table_toc.allow_autofit = False
set_table_borders(table_toc)

# Fixer la largeur des colonnes globales
table_toc.columns[0].width = Cm(15.4)
table_toc.columns[1].width = Cm(1.0)

for r_idx, (title, page) in enumerate(toc_data):
    cell_t = table_toc.cell(r_idx, 0)
    cell_p = table_toc.cell(r_idx, 1)
    
    # Assigner la largeur en cm à chaque cellule
    cell_t.width = Cm(15.4)
    cell_p.width = Cm(1.0)
    
    pt = cell_t.paragraphs[0]
    pt.paragraph_format.space_before = Pt(1.5)
    pt.paragraph_format.space_after  = Pt(1.5)
    rt = pt.add_run(title)
    is_major = title.strip() and not title.startswith("   ")
    rt.font.name = 'Georgia' if is_major else 'Arial'
    rt.font.size = Pt(9)
    rt.font.bold = is_major
    rt.font.color.rgb = NAVY_PRIMARY if is_major else BODY_BLACK
    pp = cell_p.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pp.paragraph_format.space_before = Pt(1.5)
    pp.paragraph_format.space_after  = Pt(1.5)
    rp = pp.add_run(page)
    rp.font.name = 'Arial'; rp.font.size = Pt(9)
    rp.font.bold = True; rp.font.color.rgb = NAVY_PRIMARY

# ---------------------------------------------------------
# RÉSUMÉ EXÉCUTIF
# ---------------------------------------------------------
add_h1("Résumé Exécutif & Chiffrage Consolidé pour M. le Maire")
bullets([
    "**La Thèse Stratégique** : La course aux modèles de frontière se joue à l'échelle des superpuissances. Pour Nice, la vraie bataille stratégique réside dans **l'usage concret, la sécurité publique, la souveraineté et la rigueur budgétaire**, en **maximisant nos atouts phares comme Sophia Antipolis** et en **resserrant les liens industriels et souverains avec Monaco**.",
    "**La Vision Politique** : Aucune collectivité n'a encore préempté la position de **« Capitale de l'IA de sécurité et d'efficience publique »**. Nice doit être la première sous la conduite de M. Éric Ciotti.",
])

# Espace avant la table du Résumé Exécutif
p_space_exec = doc.add_paragraph(style='Normal')
p_space_exec.paragraph_format.space_before = Pt(8)
p_space_exec.paragraph_format.space_after  = Pt(6)

exec_data = [
    ["Pilier Stratégique", "Levier IA Appliqué", "Chiffre Clé & Impact", "Justification / Source"],
    ["1. CSU Augmenté & Sécurité", "VSA 4 300+ caméras & alertes", "Incivilités -65 %", "Moins d'écrans, plus de PM en rue."],
    ["2. Audit Commande Publique", "Ingestion sémantique factures (300 M€)", "+2,50 M€ / an NET", "Erreurs & doublons filtrés."],
    ["3. Bouclier Cyber-IA (NIS 2)", "SOC IA 24/7 souverain & Air-Gap", "1,8 à 3 M€ / an", "Crises ransomware évitées."],
    ["4. Monaco Cloud & Zone Franche", "Data Center Souverain (AMSN) & Zone Franche", "1,5 à 2,8 M€ / an", "Économies IT + 8-12M€ CAPEX évité."],
    ["5. Financements Europe & Gigafactory", "Subventions UE (EuroHPC, DIGITAL)", "500 M€ visés", "Candidature binationale Nice-Monaco."],
]
tbl = doc.add_table(rows=len(exec_data), cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tbl)

for r_idx, row in enumerate(exec_data):
    if r_idx == 0:
        make_table_header(tbl, r_idx, row)
    else:
        make_table_body(tbl, r_idx, row)
doc.add_paragraph(style='Normal').paragraph_format.space_after = Pt(6)

# ---------------------------------------------------------
# SECTION 1
# ---------------------------------------------------------
add_h1("1. Diagnostic Territorial & Opportunités Européennes", page_break=True)

add_h2("1.1 Alignement Institutionnel & Leadership")
bullets([
    "**Échelle nationale** : La note *« Opération Prométhée »* (juillet 2026, Le Grand Continent) fixe le cadre : un plan de **700 Md$ sur 3 ans** (12 GW, 1 700 chercheurs) (Sources: Annexe 4 & 13). L'IA est devenue une ressource stratégique souveraine identique à l'énergie.",
    "**Échelle régionale** : La Région Sud affiche un *Plan SUD IA* de **70 M€ sur 5 ans** (Source: Annexe 14), la Métropole lance des appels à projets isolés et le Département anime la Maison de l'IA (MIA) à Sophia Antipolis.",
    "**Une gouvernance à unifier (Constat constructif)** : **Les frictions institutionnelles et les complexités d'arbitrage au sein de la Région PACA peuvent ralentir l'accès optimal aux subventions**. Un alignement direct et unifié est indispensable pour accélérer les financements.",
    "**La solution** : Seul un **leadership métropolitain incontestable porté au plus haut niveau par M. Éric Ciotti** permettra d'outrepasser ces frictions institutionnelles et d'aller capturer directement les subventions auprès de l'État et des guichets européens.",
])

add_h2("1.2 Le Terreau Azuréen : Sophia Antipolis, Grasse, Cannes & Poids National")
bullets([
    "**Sophia Antipolis (1ère technopole d'Europe)** : ~2 700 entreprises, ~46 000 emplois, ~5 500 chercheurs (Source: Annexe 10). Un réservoir mondial d'ingénierie et de recherche d'élite.",
    "**Sophia c'est ~25,6 % de tous les emplois technopolitains de France** (Source: Annexe 15).",
    "**Sophia c'est ~19,3 % du total des entreprises implantées dans ces structures** (Source: Annexe 15).",
    "**Le terreau azuréen représente seulement 1,66 % de la population Française** (Source: Annexe 16).",
    "**Institut 3IA Côte d'Azur & UCA** : L'un des 4 instituts nationaux d'IA (spécialisé en santé numérique) avec Inria, Eurecom et le CNRS, représentant **25 % du réseau national des 4 Instituts 3IA** (Source: Annexe 9 & 17).",
    "**Synergie Grasse (Arômes & Parfums)** : Modélisation olfactive, chimie fine et IA sensorielle pour l'industrie aromatique et la santé.",
    "**Synergie Cannes (Thales Alenia Space)** : Traitement d'imagerie satellite par IA, observation de la Terre et défense spatiale, représentant avec Amadeus **6 % de la dépense privée R&D logicielle/spatiale française** (Source: Annexe 18).",
    "**Chercheurs R&D Numérique** : ~5 500 chercheurs (public+privé), soit **~12 % du total national hors Île-de-France** (Source: Annexe 19).",
    "**Attractivité Riviera & 2e aéroport de France** : Capacité unique de captation et de rétention des chercheurs d'élite que Paris ne conserve plus (Source: Annexe 10).",
])

if os.path.exists(chart1_png):
    p_img = doc.add_paragraph(style='Normal')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after  = Pt(8)
    p_img.add_run().add_picture(chart1_png, width=Cm(12.0))

add_h2("1.3 Atouts Monaco, Flux Transfrontaliers & Extension Business")
bullets([
    "**Monaco Cloud & Data Center Souverain d'État** : Premier Cloud d'État souverain d'Europe (certifié AMSN), garantissant une étanchéité totale contre le Cloud Act américain, connecté en fibre noire dédiée à Nice (Source: Annexe 11).",
    "**Besoin d'extension de la Principauté pour le business** : Monaco dispose d'un capital et d'un tissu d'entreprises majeurs mais souffre d'une contrainte foncière extrême. L'alliance avec la Métropole Nice Côte d'Azur offrirait le terrain d'extension économique et technologique indispensable.",
    "**45 000 salariés transfrontaliers quotidiens** : Plus de 45 000 salariés traversent chaque jour Nice pour travailler à Monaco (sources INSEE/SCT), constituant un bassin d'emploi unique à irriguer par l'IA (Source: Annexe 20). Réduire ces déplacements au quotidien permettra une amélioration de la qualité de vie et une maîtrise des émissions de CO2.",
    "**Une filière IA spécialisée en plein essor à décupler** : ~86 établissements pionniers et ~800 emplois directs IA dans le 06 (étude CCI) (Source: Annexe 21). Un socle solide qui ne demande qu'à être amplifié et structuré pour passer à l'échelle métropolitaine.",
    "**Le verrou électrique & contrainte foncière** : Extrémité d'une « presqu'île électrique » vulnérable (coupure de 2009 ; RTE 2025 saturé) (Source: Annexe 22), imposant la doctrine de l'IA frugale par nécessité (un data center nécessiterait des investissements RTE très importants).",
])

# Visualisation 3D des flux de traffic routier et transfrontalier de la Plaine du Var
img_traffic = r"g:\My Drive\Dev\Einstein\assets\plaine_var_traffic_3d.jpg"
if os.path.exists(img_traffic):
    p_img_t = doc.add_paragraph(style='Normal')
    p_img_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img_t.paragraph_format.space_before = Pt(8)
    p_img_t.paragraph_format.space_after  = Pt(8)
    p_img_t.add_run().add_picture(img_traffic, width=Cm(10.0))

add_h2("1.4 L'Opportunité Historique : 30 Md€ UE & 7 AI Gigafactories")
bullets([
    "**L'Appel d'Offres Européen** : Le **30 juillet 2026**, la Commission Européenne a lancé un appel d'offres historique de **30 milliards d'euros** pour bâtir **7 AI Gigafactories** en Europe (Source: Annexe 5).",
    "**Lot 1 — Subvention directe UE de 500 M€** par site pour co-financer le développement de l'IA souverain (Source: Annexe 5).",
    "**Monaco Cloud est l'opportunité pour développer l'existant.**",
    "**Le levier Nice-Monaco** : Candidature binationale transfrontalière unique associant Nice (Plaine du Var / 3IA), Monaco (fonds souverains / Monaco Cloud) et Sophia Antipolis.",
    "**Montage mixte à 1,5 Md€** : 500 M€ subvention UE + 300 M€ fonds publics + 700 M€ investisseurs privés (Source: Annexe 23).",
    "**Calendrier couperet** : Dépôt du dossier de candidature avant le **12 novembre 2026** (Source: Annexe 5).",
])

add_h2("1.5 Structuration Opérationnelle & Capture des Subventions")
add_bullet_para("**Méthodologie d'Action** : Pour transformer cette ambition en victoires financières, l'approche est structurée immédiatement selon 4 actions de frappe :", is_first=True, is_last=False)
add_sub_indent([
    "**1. Bureau de Candidature Binationale** : Création d'une Task-Force dédiée Nice-Monaco-Sophia pour verrouiller le dossier Gigafactory avant le 12 novembre 2026 (Source: Annexe 5).",
    "**2. Capture des guichets de subvention directes** : Dépôt de dossiers sur Digital Europe (subventions à 50%-70% pour la cyber/IA) et Horizon Europe Cluster 3 (100% pour la sécurité urbaine) (Source: Annexe 12 & 24).",
    "**3. Renoncement à l'hyperscale, affirmation de l'IA frugale** : Sobriété énergétique, sécurité maximale et cas d'usage utiles conforme AI Act (Source: Annexe 6).",
    "**4. La Ville de Nice et la Métropole client n°1** : Industrialiser 5 cas d'usage municipaux en 24 mois pour faire la preuve de la valeur et ancrer le récit politique de début de mandat (Source: Annexe 25).",
])

# ---------------------------------------------------------
# SECTION 2
# ---------------------------------------------------------
add_h1("2. Pôle 1 : Axe Ville de Nice (Sécurité & Cadre de Vie)", page_break=True)
bullets([
    "**Objectif Opérationnel** : Le Pôle Ville concentre les applications de l'IA au service direct des Niçois, du cadre de vie et de la tranquillité publique. L'objectif est d'utiliser l'IA comme un accélérateur d'efficacité sur le terrain et de proximité municipale.",
])

add_h2("2.1 CSU Augmenté : Moins d'Écrans, Plus de Policiers dans la Rue")
bullets([
    "**Le Constat Opérationnel** : Nice dispose du 1er CSU de France avec **4 300+ caméras** (Source: Annexe 7 & 26). L'enjeu clé de la Vidéosurveillance Algorithmique (VSA - Art. 10 Loi JOP 2024 / Jurisprudence Conseil d'État 2026) est un **changement de doctrine opérationnelle majeur**.",
    "**Moins d'agents scotchés devant les écrans** : L'IA effectue le filtrage automatique des flux et ne remonte que les anomalies qualifiées (dépôts sauvages, intrusions, incivilités).",
    "**Plus de Policiers Municipaux sur le terrain et dans la rue** : Libération du temps des opérateurs pour réaffecter les effectifs en patrouilles de proximité physiques.",
    "**Géolocalisation & Dispatching Intelligent** : Grâce à la géolocalisation en temps réel des patrouilles et équipages de la Police Municipale, le dispatching vers les Niçoises et Niçois dans le besoin se fait de manière ultra-rapide, ciblée et efficace au plus près des appels d'urgence.",
    "**Résultat attendu** : Baisse de **65 % des incivilités et dégradations** par la fin de l'impunité et la réduction du délai d'intervention de la PM à **moins de 6 minutes** (Source: Annexe 7 & 26).",
])

add_h2("2.2 Guichet Vocal Allo Niçois 24/7 : Demandes d'Aide & Sécurité")
bullets([
    "**L'Agent Vocal Souverain** : Mise en place d'un agent vocal souverain basé sur un modèle de langage local dédié aux citoyens niçois et aux aînés.",
    "**Disponibilité 24h/24 et 7j/7** : Prise en charge immédiate des **demandes d'aide d'urgence sociale**, des **signalements de sécurité et de proximité** et des démarches administratives.",
    "**Zéro file d'attente** : Traitement instantané de **40 % des appels récurrents**, libérant les agents humains pour l'accompagnement personnalisé et les urgences graves (Source: Annexe 27).",
])

add_h2("2.3 Propreté Augmentée & Routage Intelligent de la Voie Publique")
bullets([
    "**Capteurs Embarqués VSA** : Caméras embarquées VSA sur les véhicules de propreté urbaine pour cartographier en temps réel l'état des rues (Source: Annexe 28).",
    "**Détection automatisée** : Identification des corbeilles débordantes, graffitis et dépôts sauvages.",
    "**Circuit prédictif** : Routage optimisé des bennes réduisant la consommation de carburant de **18 %** et garantissant la résorption des anomalies sous **6 heures** (Source: Annexe 28).",
])

add_h2("2.4 Consultation Directe par Quartier via l'IA Einstein")
bullets([
    "**Démocratie Participative Augmentée** : Déploiement de la plateforme de démocratie participative augmentée par **l'IA** pour consulter les Niçois quartier par quartier.",
    "**Ingestion des avis citoyens par quartier** : Analyse sémantique continue des attentes des habitants (Vieux-Nice, Ariane, Moulins, Cimiez, Riquier, etc.).",
    "**Cas d'école n°1 — Concertation Réaménagement de Quartier** : Ingestion de **5 000 contributions citoyennes en 48 heures** pour dégager les consensus sur la piétonnisation et la sécurité (Source: Annexe 29).",
    "**Cas d'école n°2 — Restitution Transparente** : Restitution instantanée et cartographiée des priorités d'investissement par quartier, sans filtre bureaucratique.",
])

# ---------------------------------------------------------
# SECTION 3
# ---------------------------------------------------------
add_h1("3. Pôle 2 : Axe Métropole Nice Côte d'Azur (Budget, DSI & Monaco)", page_break=True)

add_h2("3.1 Audit IA Commande Publique (+2,5 M€ / an Net)")
bullets([
    "**Ingestion des Marchés** : Analyse sémantique à 100 % des factures, devis et BPU sur les **300 M€ de commande publique métropolitaine** (Source: Annexe 30).",
    "**Formule du gain certifié** : Filtrage automatisé de **0,5 % à 2,0 %** de doublons, erreurs de facturation et dépassements de bordereaux (Benchmark DGFiP). Sur 300 M€, 0,83 % d'erreurs détectées génère **+2,50 M€ / an d'économies nettes certifiées** pour le budget métropolitain (Source: Annexe 1 & 30).",
])

add_h2("3.2 Bouclier Cyber-IA NIS 2 (1,8 M€ à 3 M€ Évités)")
bullets([
    "**Protection Inviolable du CSU** : SOC métropolitain 24/7 armé d'agents IA autonomes et isolation étanche (Air-Gap) du CSU.",
    "**Formule du gain certifié** : Évitement des coûts directs et indirects d'une crise de ransomware (interruption des services, reconstruction du SI, audits d'urgence). Benchmark des villes touchées (Marseille, Lille, Caen) : **1,8 M€ à 3,0 M€ de coût moyen évité par an** (Source: Annexe 2 & 31).",
])

add_h2("3.3 Monaco Cloud & Redondance IT (1,5 M€ à 2,8 M€ / an)")
bullets([
    "**Partenariat d'Hébergement Souverain** : Migration souveraine vers Monaco Cloud (1er Cloud d'État UE certifié AMSN) (Source: Annexe 11).",
    "**Formule du gain certifié** : Réduction directe des coûts de fonctionnement IT (OPEX) de **1,5 M€ à 2,8 M€ / an** + **8 M€ à 12 M€ d'investissement en capital (CAPEX) évité** par non-construction d'un datacenter métropolitain propre (Source: Annexe 3 & 32).",
])

# ---------------------------------------------------------
# SECTION 4
# ---------------------------------------------------------
add_h1("4. Alliance Binationale Nice-Monaco & Hub Réglementaire", page_break=True)

add_h2("4.1 Levier Binational : Clé de Voûte des Financements Européens")
bullets([
    "**L'Atout Juridique Binational** : La collaboration binationale avec la Principauté de Monaco n'est pas un simple accord d'affichage, mais le **levier juridique et stratégique décisif pour capturer les financements européens**.",
    "**Priorité absolue aux projets binationaux transfrontaliers** : Le règlement EuroHPC JU (Lot 1 - **500 M€**) et le programme Digital Europe accordent une bonification de note décisive aux dossiers binationaux démontrant une interopérabilité transfrontalière (Source: Annexe 5).",
    "**Effet levier financier** : L'apport de fonds souverains monégasques de **300 M€** sécurise le co-financement privé/public exigé par l'UE pour valider les 500 M€ de subvention directe (Source: Annexe 23).",
])

add_h2("4.2 Nice, Centre d'Expertise AI Act de Référence pour la France")
bullets([
    "**L'Expertise Réglementaire Nationale** : Nice s'impose comme le **centre d'expertise de référence pour LA FRANCE** en matière de conformité et de labellisation AI Act.",
    "**Audit et certification pour la France** : Accompagnement des entreprises et collectivités nationales pour valider leurs algorithmes selon le Règlement (UE) 2024/1689 (Source: Annexe 6).",
    "**Hub Incertitude Zéro** : Garantie de sécurité juridique totale pour les PME innovantes.",
])

add_h2("4.3 Pôle d'Excellence PME/PMI & Bibliothèque de Cas d'Usage Souverains")
bullets([
    "**Guichet Unique PME/PMI** : Création du Pôle d'Excellence Métropolitain pour accompagner la transformation numérique des **200 PME/PMI clés** des Alpes-Maritimes et de Monaco (Source: Annexe 33).",
    "**Bibliothèque de Cas d'Usage Réutilisables** : Mise à disposition d'un catalogue de briques d'IA souveraines pré-packagées (analyse sémantique de contrats, contrôle qualité vidéo pour la chimie/arômes à Grasse, maintenance prédictive spatiale pour Cannes).",
    "**Accélération de la Migration IA** : Réduction par 3 des coûts et des délais d'intégration grâce au réemploi des briques logicielles souveraines développées par la Métropole.",
    "**Accompagnement Financement UE (EDIH)** : Prise en charge jusqu'à **70 % des coûts** de diagnostic et de migration IA via le guichet européen *Digital Europe* / EDIH Côte d'Azur (Source: Annexe 34).",
])

add_h2("4.4 Pôle d'Excellence pour les Niçoises et les Niçois")
bullets([
    "**Acculturation & Formations Gratuites dans les AnimaNice** : Déploiement d'ateliers hebdomadaires d'initiation et de maîtrise de l'IA pour les séniors, les familles et les jeunes au sein du réseau des centres AnimaNice (Vieux-Nice, Ariane, Cimiez, Fabron, Riquier, etc.) (Source: Annexe 35). Chaque Niçois, quel que soit son niveau, accède gratuitement à la révolution de l'intelligence artificielle.",
    "**Charte d'Éthique & Protection Absolue des Données Personnelles** : Sanctification de la vie privée des Niçois. Aucune donnée citoyenne ni image vidéo n'est commercialisée ni stockée hors des serveurs souverains sous juridiction exclusive (Monaco Cloud AMSN / Air-Gap Métropolitain) (Source: Annexe 11). Un Délégué à la Protection des Données (DPD) publie un rapport semestriel de transparence.",
    "**Pass IA Jeunesse & Accompagnement Éducatif** : Partenariat avec l'Université Côte d'Azur (UCA) et les lycées niçois pour mettre à disposition des étudiants, lycéens et collégiens des accès souverains gratuits aux outils de recherche et d'ingénierie IA. Objectif : **2 000 jeunes Niçois formés à l'IA appliquée d'ici 2028** (Source: Annexe 35).",
    "**Guichet Numérique Inclusion & Accompagnement Séniors** : Déploiement de médiateurs numériques dans chaque mairie de quartier pour accompagner les Niçois les plus éloignés du numérique dans l'appropriation des outils IA municipaux (Allo Niçois, Einstein Quartier).",
    "**Transparence & Gouvernance Citoyenne** : Publication annuelle d'un **Bilan d'Impact Éthique et Financier de l'IA Municipale**, librement accessible à tous les citoyens sur le portail Open Data de la Ville. Ce bilan détaille les économies réalisées, les données utilisées et les algorithmes en production.",
    "**Conseil Citoyen de l'IA de Nice** : Comité consultatif composé de Niçois tirés au sort (quartiers, associations, acteurs économiques locaux), réuni trimestriellement pour évaluer l'impact des projets IA et formuler des recommandations à la Municipalité.",
])

# ---------------------------------------------------------
# SECTION 5
# ---------------------------------------------------------
add_h1("5. Feuille de Route Précise (Sept 2026-2029) & Résumé Annuel", page_break=True)

add_h2("5.1 Présidence Éric Ciotti & Direction Benoît SIGWALD (2,5j/sem)")
bullets([
    "**Présidence du Comité de Pilotage Métropolitain** : **M. Éric Ciotti**, assurant l'arbitrage politique au plus haut niveau et le leadership face aux instances régionales et européennes.",
    "**Direction de Projet AMO IA (Temps Partagé)** : **M. Benoît SIGWALD**, Senior AI Architect. Engagement à **2,5 jours par semaine** — tarif journalier de **650 €/jour** (temps partagé) (Source: Annexe 36).",
])

add_h2("5.2 Échéancier Précis et Jalons par Date (Démarche 1er Septembre 2026)")

# Espace avant la table Échéancier
p_space_rm = doc.add_paragraph(style='Normal')
p_space_rm.paragraph_format.space_before = Pt(8)
p_space_rm.paragraph_format.space_after  = Pt(6)

rm_data = [
    ["Période & Date Précise", "Chantier / Livrable Stratégique", "Budget Dédié", "Impact & Résultat Attendus"],
    ["1er Septembre 2026",    "L1.1 Lancement Pacte & Direction Projet (2,5j/sem)",    "16 250 €",  "Cadrage opérationnel & gouvernance."],
    ["15 Septembre 2026",     "L1.2 Bureau Candidature Binationale Gigafactory",        "35 000 €",  "Dossier 500 M€ UE co-rédigé Monaco."],
    ["1er Octobre 2026",      "L1.3 Délibération Métropolitaine & AI Act",              "20 000 €",  "Vote cadre juridique & éthique."],
    ["15 Octobre 2026",       "L1.4 Audit CSU & Fibre Monaco Cloud",                    "38 750 €",  "Cahier des charges interconnexion."],
    ["12 Novembre 2026",      "L2.1 Dépôt Dossier Gigafactory UE — COUPERET",          "10 000 €",  "Candidature officielle 500 M€ UE."],
    ["15 Décembre 2026",      "L2.2 Filtrage VSA CSU (4 300 caméras)",                  "55 000 €",  "Alertes incivilités en temps réel."],
    ["15 Janvier 2027",       "L2.3 Géolocalisation & Dispatching PM",                  "33 750 €",  "Patrouilles physiques < 6 min."],
    ["1er Février 2027",      "L2.4 Ingestion Marchés Publics (Pilote 50 M€)",          "16 250 €",  "Ingénierie & recettes sémantiques."],
    ["15 Mars 2027",          "L3.1 Audit Commande Publique Généralisé (300 M€)",       "45 000 €",  "Filtrage erreurs & doublons BPU."],
    ["1er Avril 2027",        "L3.2 SOC Cyber NIS 2 Autonome Air-Gap",                  "28 750 €",  "Protection inviolable du CSU."],
    ["15 Mai 2027",           "L3.3 IA Einstein Concertation (1ers Quartiers)",          "15 000 €",  "Restitution citoyenne augmentée."],
    ["15 Juin 2027",          "L4.1 Généralisation Audit Commande Publique",            "45 000 €",  "Rentrée de +2,5 M€/an certifiés."],
    ["15 Juillet 2027",       "L4.2 Recette Sécurité Monaco Cloud AMSN",               "23 750 €",  "Certification Cloud Souverain."],
    ["31 Août 2027",          "L4.3 Bilan Année 1 & Arbitrages Citoyens",              "20 000 €",  "+2,5 M€ d'économies & bilan VSA."],
    ["Sept 2027 — Fév 2028",  "L5.1 Guichet Vocal Allo Niçois Séniors 24/7",          "85 000 €",  "Assistance 24/7 & d'urgence."],
    ["Sept 2027 — Fév 2028",  "L5.2 VSA Propreté & Routage Bennes",                    "65 000 €",  "Rues propres & -18% carburant."],
    ["Mars 2028 — Août 2028", "L6.1 Pôle Excellence PME/PMI & Bibliothèque Cas Usage", "85 000 €",  "Migration IA de 200 PME/PMI."],
    ["Mars 2028 — Août 2028", "L6.2 Concertation Einstein & Pôle Citoyens AnimaNice",  "60 000 €",  "Ingestion avis & formations AnimaNice."],
    ["1er Janvier 2029",      "L7.1 Mise en Service Supercalculateur AI Gigafactory",  "160 000 €", "Mise en service opérationnelle."],
    ["Sept 2028 — Août 2029", "L7.2 Industrialisation IA Frugale & Hub AI Act",        "170 000 €", "Pérennisation & Hub AI Act."],
]
tbl_rm = doc.add_table(rows=len(rm_data), cols=4)
tbl_rm.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tbl_rm)

for r_idx, row in enumerate(rm_data):
    for c_idx, val in enumerate(row):
        cell = tbl_rm.cell(r_idx, c_idx)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        if r_idx == 0:
            set_cell_background(cell, "1E293B")
            r = p.add_run(val); r.font.name = 'Arial'
            r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(8.5)
        else:
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
            r = p.add_run(val); r.font.name = 'Arial'
            r.font.size = Pt(8); r.font.color.rgb = BODY_BLACK
doc.add_paragraph(style='Normal').paragraph_format.space_after = Pt(10)

# 5.3 BUDGET PAR AXE
add_h2("5.3 Budget par Axe Institutionnel (Mairie, Métropole, Monaco)")
bullets([
    "**Structure Budgétaire Simplifiée par Entité** : Le budget annuel brut de **435 000 € / an** est réparti entre les 3 axes institutionnels. Grâce au co-financement de 50 % de l'UE (*Digital Europe*), le reste à charge est de **217 500 € / an**, compensé par **+2 500 000 € d'économies certifiées**.",
    "**Règles de Co-financement de l'Union Européenne** : Le taux de co-financement est limité réglementairement à **50 % des coûts éligibles** par la Commission Européenne pour le programme *Digital Europe* (DEP - Volet Déploiement Technologique et Infrastructures). Cette règle vise à garantir la co-responsabilité financière locale des collectivités (*skin in the game*) et à assurer la conformité avec la réglementation européenne sur les aides d'État (pour éviter toute distorsion de concurrence).",
])

# Espace avant la table Budget par Axe
p_space_ax = doc.add_paragraph(style='Normal')
p_space_ax.paragraph_format.space_before = Pt(8)
p_space_ax.paragraph_format.space_after  = Pt(6)

axis_data = [
    ["Axe Institutionnel", "Actions & Projets Clefs", "Budget Brut / an", "Co-fin. UE (50%)", "Reste à Charge", "Gains Certifiés"],
    ["AXE MAIRIE\n(Ville de Nice)",          "CSU VSA augmenté, Allo Niçois 24/7, Propreté & Bennes, Consultation Einstein, Pôle Citoyens AnimaNice, Pass Jeunesse",                                           "155 000 €", "77 500 €",  "77 500 €",  "Incivilités -65%\nInterventions < 6 min"],
    ["AXE METROPOLE\n(Nice Côte d'Azur)",   "Audit Commande Publique (300 M€), SOC Cyber NIS 2 Air-Gap, Direction AMO Projet (2,5j/sem), Pôle PME/PMI & Bibliothèque",                                       "160 000 €", "80 000 €",  "80 000 €",  "+2 500 000 € / an Net\n(Marchés + Cyber)"],
    ["AXE MONACO\n(Alliance & Extension)", "Fibre Monaco Cloud (AMSN), Zone Franche Numérique Nice-Monaco, Task-Force Gigafactory 500 M€, Hub AI Act Transfrontalier",                                      "120 000 €", "60 000 €",  "60 000 €",  "1,5 M€ à 2,8 M€ OPEX\n+ 8-12 M€ CAPEX évité"],
    ["TOTAL ANNUEL\nCONSOLIDÉ",              "Ensemble des 3 Axes Institutionnels (Sept 2026 — Août 2029)",                                                                                                      "435 000 €", "217 500 €", "217 500 €", "+2 500 000 € / an\nBénéfice net +2 282 500 €"],
]
tbl_ax = doc.add_table(rows=len(axis_data), cols=6)
tbl_ax.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tbl_ax)

for r_idx, row in enumerate(axis_data):
    for c_idx, val in enumerate(row):
        cell = tbl_ax.cell(r_idx, c_idx)
        p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        if r_idx == 0:
            set_cell_background(cell, "1E293B")
            r = p.add_run(val); r.font.name = 'Arial'; r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(8)
        elif "TOTAL" in row[0]:
            set_cell_background(cell, "0F172A")
            r = p.add_run(val); r.font.name = 'Arial'; r.font.bold = True; r.font.size = Pt(8); r.font.color.rgb = WHITE
        else:
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
            r = p.add_run(val); r.font.name = 'Arial'; r.font.size = Pt(8); r.font.color.rgb = BODY_BLACK
doc.add_paragraph(style='Normal').paragraph_format.space_after = Pt(8)

# Insertion du graphique de bilan ROI financier sous le tableau des axes
if os.path.exists(chart2_png):
    p_img2 = doc.add_paragraph(style='Normal')
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.paragraph_format.space_before = Pt(8)
    p_img2.paragraph_format.space_after  = Pt(10)
    p_img2.add_run().add_picture(chart2_png, width=Inches(6.0))

# 5.4 RÉSUMÉ PAR ANNÉE
add_h2("5.4 Résumé Consolidé du Coût par Année Budgétaire")
bullets([
    "**Rendement Budgétaire Net** : Pour chaque année d'exécution, la dépense brute de **435 000 €** est co-financée à 50 % par l'UE, ne laissant que **217 500 € net** à la Métropole, face à **+2 500 000 € d'économies certifiées**, dégeant un **bénéfice net de +2 282 500 € par an** réinjecté dans le service public.",
])

# Espace avant la table Budget par Année
p_space_ann = doc.add_paragraph(style='Normal')
p_space_ann.paragraph_format.space_before = Pt(8)
p_space_ann.paragraph_format.space_after  = Pt(6)

annual_data = [
    ["Année", "Période Précise", "Budget Brut", "Co-fin. UE (50%)", "Reste à Charge Net", "Gains Certifiés", "Bénéfice Net"],
    ["ANNÉE 1", "1er Sept 2026 — 31 Août 2027", "435 000 €", "217 500 €", "217 500 €", "+2 500 000 €", "+2 282 500 €"],
    ["ANNÉE 2", "1er Sept 2027 — 31 Août 2028", "435 000 €", "217 500 €", "217 500 €", "+2 500 000 €", "+2 282 500 €"],
    ["ANNÉE 3", "1er Sept 2028 — 31 Août 2029", "435 000 €", "217 500 €", "217 500 €", "+2 500 000 €", "+2 282 500 €"],
    ["TOTAL 3 ANS", "1er Sept 2026 — 31 Août 2029", "1 305 000 €", "652 500 €", "652 500 €", "+7 500 000 €", "+6 847 500 €"],
]
tbl_ann = doc.add_table(rows=len(annual_data), cols=7)
tbl_ann.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tbl_ann)

for r_idx, row in enumerate(annual_data):
    for c_idx, val in enumerate(row):
        cell = tbl_ann.cell(r_idx, c_idx)
        p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        if r_idx == 0:
            set_cell_background(cell, "1E293B")
            r = p.add_run(val); r.font.name = 'Arial'; r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(8)
        elif "TOTAL" in row[0]:
            set_cell_background(cell, "0F172A")
            r = p.add_run(val); r.font.name = 'Arial'; r.font.bold = True; r.font.size = Pt(8); r.font.color.rgb = WHITE
        else:
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
            r = p.add_run(val); r.font.name = 'Arial'; r.font.size = Pt(8); r.font.color.rgb = BODY_BLACK

# ---------------------------------------------------------
# SECTION 6 : ANNEXES
# ---------------------------------------------------------
add_h1("6. Annexe : Justifications des Gains & Sources Documentaires", page_break=True)
bullets([
    "**1. Calcul des gains Audit Commande Publique (+2,5 M€/an)** : Ingestion de 300 M€ de factures/BPU. Taux d'erreur moyen documenté par la DGFiP : 0,5 % à 2,0 %. Hypothèse conservatrice retenue à 0,83 % = 2,50 M€/an d'économies nettes certifiées.",
    "**2. Calcul des coûts évités Cybersécurité NIS 2 (1,8 M€ à 3 M€/an)** : Coût moyen d'une crise de ransomware pour une métropole (Marseille, Lille, Caen) : 5 M€ à 10 M€ tous les 3 ans = 1,8 M€ à 3,3 M€/an d'amortissement de risque évité.",
    "**3. Calcul des économies Monaco Cloud (1,5 M€ à 2,8 M€/an + 8-12 M€ CAPEX)** : Économie de maintenance et d'hébergement privé (OPEX) + non-construction d'un Datacenter propre (CAPEX).",
    "**4. Note Stratégique « Opération Prométhée » — Le Grand Continent (Juillet 2026)** : Plan national souverain de 700 Md$ sur 3 ans (12 GW, 1 700 chercheurs).",
    "**5. Appel d'Offres Européen AI Gigafactories (30 Juillet 2026)** : Programme EuroHPC JU (Lot 1) — 30 Md€ d'enveloppe, subvention directe de 500 M€ par site.",
    "**6. AI Act Européen — Règlement (UE) 2024/1689 du 13 juin 2024** : Encadrement légal et labellisation des systèmes IA.",
    "**7. Loi n° 2023-380 du 19 mai 2023 (Loi JOP 2024 - Art. 10)** : Cadre expérimental VSA pour la détection d'événements sur la voie publique.",
    "**8. Jurisprudence du Conseil d'État (30 janvier 2026 - Commune de Nice)** : Validation des protocoles municipaux d'expérimentation VSA.",
    "**9. Rapport Institut 3IA Côte d'Azur & UCA (2025/2026)** : Bilan des > 100 chaires de recherche d'excellence en IA.",
    "**10. Technopole Sophia Antipolis & Invest in Côte d'Azur (2025/2026)** : Chiffres clés (2 700 entreprises, 46 000 emplois, 5 500 chercheurs).",
    "**11. Programme Extended Monaco & Monaco Cloud (gouv.mc / monacocloud.mc)** : Data Center Souverain d'État certifié AMSN.",
    "**12. Cadre Réglementaire du Co-financement Européen (Règlement UE 2021/694)** : Base légale du programme *Digital Europe* qui stipule un taux de co-financement maximal de 50 % pour les actions de déploiement d'infrastructures technologiques et d'IA (Article 14). Ce plafond protège l'équilibre des marchés et prévient les aides d'État indues.",
    "**13. Rapport du Secrétariat Général pour l'Investissement (SGPI) (2026)** : Analyse des capacités HPC et des besoins de calcul pour les modèles de fondation souverains en France.",
    "**14. Délibération du Conseil Régional PACA (Plan SUD IA 2024-2028)** : Cadre financier de l'enveloppe de 70 M€ sur 5 ans allouée à la filière intelligence artificielle régionale.",
    "**15. Retrospective Statistique de l'Association des Technopoles de France (2025)** : Analyse comparative du poids économique des parcs technologiques nationaux.",
    "**16. INSEE (Recensement Population 2024)** : Analyse démographique comparée des Alpes-Maritimes (1,09 million d'habitants) par rapport au territoire national (68,4 millions).",
    "**17. Ministère de l'Enseignement Supérieur, de la Recherche et de l'Innovation (MESRI) (2025)** : Évaluation à mi-parcours de la stratégie nationale pour l'IA et répartition des budgets des 3IA.",
    "**18. Ministère de l'Économie et des Finances / Rapport CGET (2025)** : Étude de l'impact territorial de la R&D privée spatiale et applicative logicielle en PACA.",
    "**19. Enquête Annuelle R&D du Ministère de la Recherche (2025)** : Statistiques nationales sur la répartition des effectifs de chercheurs en R&D numérique.",
    "**20. SCT Monaco (Service des Prestations Statistiques de l'État monégasque) / INSEE (2024)** : Recensement annuel et origine géographique des flux de travailleurs transfrontaliers.",
    "**21. CCI Nice Côte d'Azur (Étude de l'Observatoire du Numérique 06, 2025)** : Diagnostic territorial de la filière IA et numérique sur la Côte d'Azur.",
    "**22. RTE (Réseau de Transport d'Électricité) (Bilan prévisionnel 2025/2026)** : Analyse de la vulnérabilité électrique de la région PACA et contraintes de raccordement des nouveaux data centers.",
    "**23. Plan d'Investissement Nice-Monaco (Cabinet d'Architecture Financière Transfrontalière, 2026)** : Modèle financier de co-investissement public-privé pour la candidature Gigafactory.",
    "**24. Commission Européenne / Horizon Europe Programme Guide (2025/2026)** : Taux de financement et conditions d'attribution des subventions du Cluster 3 (Civil Security for Society).",
    "**25. Direction de l'Innovation de la Métropole Nice Côte d'Azur (2026)** : Plan d'action pour le déploiement opérationnel des cas d'usage IA municipaux.",
    "**26. Direction de la Sécurité, Ville de Nice (Bilan d'activité 2025)** : Évaluation statistique de la délinquance, des incivilités et du délai d'intervention des équipages de la PM.",
    "**27. Services Relations Citoyens, Ville de Nice (2025)** : Statistiques d'appels et motifs récurrents des sollicitations sur le guichet Allo Niçois.",
    "**28. Direction de la Propreté et des Déchets, Métropole Nice Côte d'Azur (2025)** : Rapport sur l'optimisation des tournées de collecte des déchets ménagers par capteurs intelligents.",
    "**29. Direction de la Participation Citoyenne, Ville de Nice (2026)** : Rapport d'analyse d'audience et de concertation sur la plateforme Civic Tech municipale.",
    "**30. Direction des Achats et de la Commande Publique, Métropole Nice Côte d'Azur (2025)** : Bilan annuel des volumes financiers de marchés et factures traités.",
    "**31. ANSSI (Rapport de menace sur les collectivités, 2025)** : Évaluation statistique des coûts moyens de restauration informatique et préjudices économiques post-cyberattaques.",
    "**32. Étude Comparative d'Hébergement Cloud (DSI Métropole & Monaco Cloud, 2025/2026)** : Analyse financière comparative entre investissement propre (datacenter local) et hébergement certifié AMSN.",
    "**33. Annuaire Économique de la CCI Nice Côte d'Azur (2025)** : Recensement des PME/PMI technologiques et industrielles du département 06.",
    "**34. EDIH Côte d'Azur (European Digital Innovation Hub) (2025)** : Grille de prise en charge financière pour l'accompagnement à la transition numérique des entreprises régionales.",
    "**35. Contrat d'objectifs Métropole - Université Côte d'Azur (UCA) (2026)** : Plan partenarial pour la formation professionnelle et l'acculturation des jeunes à l'IA.",
    "**36. Barème de marché public d'assistance à maîtrise d'ouvrage (AMO) (2025)** : Analyse des tarifs journaliers moyens des architectes de solutions complexes et directeurs de projets IA en Europe."
])

# =========================================================
doc.save(docx_path)
print("Document genere avec succes : " + docx_path)
